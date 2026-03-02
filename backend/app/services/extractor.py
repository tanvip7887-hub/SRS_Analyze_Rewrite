import json
import re
from docx import Document

# -----------------------------------------
# STEP 1: Load Document
# -----------------------------------------

def load_doc(path):
    doc = Document(path)
    return doc


# -----------------------------------------
# STEP 2: Extract text from paragraphs
# -----------------------------------------

def extract_paragraph_text(doc):
    text_list = []

    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            text_list.append(t)

    return text_list


# -----------------------------------------
# STEP 3: Extract text from tables
# -----------------------------------------

def extract_table_text(doc):
    text_list = []

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    text_list.append(t)

    return text_list


# -----------------------------------------
# STEP 4: Combine both sources
# -----------------------------------------

def combine_text(paragraphs, table_text):
    combined = paragraphs + table_text
    return combined


# -----------------------------------------
# STEP 5: Cleaning Helpers
# -----------------------------------------

REQ_PATTERN = r'\b(FR|NFR|SR|DR|IR)-\d+\b'

def clean_requirement_text(text):
    text = re.sub(r'^[\.\-\•\:\s]+', '', text)
    text = re.sub(r'\bID Requirement\b', '', text, flags=re.IGNORECASE)
    text = re.sub(r'\s+', ' ', text)
    return text.strip()


def is_valid_requirement(text):
    if not text:
        return False

    if len(text.split()) < 5:
        return False

    if text.lower() in ["requirement", "id", "id requirement"]:
        return False

    return True


# -----------------------------------------
# STEP 6: Extract requirements (Improved)
# -----------------------------------------

def extract_requirements(text_blocks):
    req_dict = {}
    current_id = None

    for line in text_blocks:
        match = re.search(REQ_PATTERN, line)

        if match:
            req_id = match.group()
            text_after = line.split(req_id, 1)[-1]

            cleaned_text = clean_requirement_text(text_after)

            req_dict[req_id] = cleaned_text
            current_id = req_id

        else:
            if current_id:
                req_dict[current_id] += " " + line

    # Final Cleaning + Validation
    final_dict = {}

    for req_id, text in req_dict.items():
        cleaned_text = clean_requirement_text(text)

        if is_valid_requirement(cleaned_text):
            final_dict[req_id] = cleaned_text

    return final_dict


# -----------------------------------------
# STEP 7: Save to JSON
# -----------------------------------------

def save_json(data, path):
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=4, ensure_ascii=False)
    return True


# -----------------------------------------
# MAIN FUNCTION FOR BACKEND
# -----------------------------------------

def run_extractor(docx_path, output_json_path):
    """
    This is the MAIN FUNCTION backend will call
    → upload.py will use this function
    """

    doc = load_doc(docx_path)
    paragraphs = extract_paragraph_text(doc)
    table_text = extract_table_text(doc)
    all_blocks = combine_text(paragraphs, table_text)
    requirements = extract_requirements(all_blocks)

    save_json(requirements, output_json_path)

    return requirements  # FastAPI will return this