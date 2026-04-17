import json
from docx import Document


def annotate_srs(requirements: dict,
                 duplicate_report: dict,
                 ambiguity_report: list,
                 generate_docx: bool = False,
                 docx_path: str = "annotated_srs.docx"):
    """
    This function generates:
    - summary
    - duplicates list
    - ambiguous requirements
    - format issues
    - optional DOCX output

    Inputs:
        requirements: {"FR-01": "text", ...}
        duplicate_report: {"duplicates": [[{id,text}, {id,text}], ...]}
        ambiguity_report: [{"id": "...", "text": "...", ...}, ...]
    """

    # Normalize format
    if isinstance(requirements, list):
        requirements = {item["id"]: item["text"] for item in requirements}

    # ------------------------------
    # PROCESS DUPLICATES
    # ------------------------------
    duplicate_pairs = []
    duplicate_removed_ids = set()

    if isinstance(duplicate_report, dict) and "duplicates" in duplicate_report:
        duplicate_groups = duplicate_report["duplicates"]
    else:
        duplicate_groups = []

    for group in duplicate_groups:
        if not isinstance(group, list) or len(group) < 2:
            continue

        master = group[0]["id"]
        for dup in group[1:]:
            duplicate_pairs.append({"req1": master, "req2": dup["id"]})
            duplicate_removed_ids.add(dup["id"])

    # ------------------------------
    # PROCESS AMBIGUITY ANALYSIS
    # ------------------------------
    ambiguous_list = []

    for item in ambiguity_report:
        req_id = item.get("id")
        if not req_id or req_id in duplicate_removed_ids:
            continue

        ambiguous_list.append({
            "id": req_id,
            "text": item.get("text", ""),
            "ambiguous_flags": item.get("ambiguous_flags", []),
            "ambiguity_score": item.get("ambiguity_score", 0)
        })

    ambiguous_list.sort(key=lambda x: x["ambiguity_score"], reverse=True)

    # ------------------------------
    # FORMAT ISSUES
    # ------------------------------
    format_issues = []

    for req_id, text in requirements.items():
        if req_id in duplicate_removed_ids:
            continue

        low = text.lower()
        if "shall" not in low and "should" not in low:
            format_issues.append({
                "id": req_id,
                "text": text,
                "reason": "Missing SHALL/SHOULD keyword"
            })

    # ------------------------------
    # BUILD FINAL REPORT
    # ------------------------------
    report = {
        "summary": {
            "total_requirements": len(requirements),
            "duplicate_groups": len(duplicate_groups),
            "duplicates_removed": len(duplicate_removed_ids),
            "ambiguous_count": len(ambiguous_list),
            "format_issue_count": len(format_issues)
        },
        "duplicates": duplicate_pairs,
        "ambiguous_requirements": ambiguous_list,
        "format_issues": format_issues
    }

    # ------------------------------
    # OPTIONAL DOCX GENERATION
    # ------------------------------
    if generate_docx:
        doc = Document()
        doc.add_heading("Annotated Software Requirements Specification", level=1)

        # Summary
        doc.add_heading("1. Summary", level=2)
        for key, val in report["summary"].items():
            doc.add_paragraph(f"{key.replace('_', ' ').title()}: {val}")

        # Duplicate section
        doc.add_heading("2. Duplicate Requirements", level=2)
        if duplicate_pairs:
            for pair in duplicate_pairs:
                doc.add_paragraph(f"{pair['req1']} is duplicate of {pair['req2']}", style='List Bullet')
        else:
            doc.add_paragraph("No duplicates found.")

        # Ambiguous section
        doc.add_heading("3. Ambiguous Requirements", level=2)
        if ambiguous_list:
            for item in ambiguous_list:
                doc.add_heading(item["id"], level=3)
                doc.add_paragraph(f"Requirement: {item['text']}")
                doc.add_paragraph(f"Ambiguity Score: {item['ambiguity_score']}")
                if item["ambiguous_flags"]:
                    doc.add_paragraph("Ambiguity Flags:")
                    for flag in item["ambiguous_flags"]:
                        doc.add_paragraph(flag, style='List Bullet')
        else:
            doc.add_paragraph("No ambiguous requirements found.")

        # Format issues
        doc.add_heading("4. Format Issues", level=2)
        if format_issues:
            for item in format_issues:
                doc.add_heading(item["id"], level=3)
                doc.add_paragraph(f"Requirement: {item['text']}")
                doc.add_paragraph(f"Issue: {item['reason']}")
        else:
            doc.add_paragraph("No format issues found.")

        doc.save(docx_path)

    return report